#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
信息化项目验收文档编写技能 - Markdown转Word脚本

功能：将markdown文档转换为Word文档
作者：Claude
日期：2026-01-27

依赖：
- python-docx: 用于创建Word文档
- markdown: 用于解析markdown文本

安装依赖：
pip install python-docx markdown
"""

import os
import re
import argparse
from pathlib import Path
from typing import List, Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 缺少必要的库")
    print("请安装依赖: pip install python-docx markdown")
    exit(1)


class MarkdownToWordConverter:
    """Markdown转Word转换器"""

    def __init__(self):
        self.doc = None
        self.current_level = 0
        self.in_code_block = False
        self.in_table = False
        self.table_data = []
        self.table_headers = []

    def create_word_document(self) -> Document:
        """创建Word文档并设置样式"""
        doc = Document()

        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置标题样式
        self._setup_heading_styles(doc)

        return doc

    def _setup_heading_styles(self, doc: Document):
        """设置标题样式"""
        # 检查样式是否已存在，如果存在则跳过
        try:
            # 标题1
            if 'Heading 1' not in doc.styles:
                style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = '黑体'
                style.font.size = Pt(16)
                style.font.bold = True
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                style.paragraph_format.space_after = Pt(12)
        except:
            pass

        try:
            # 标题2
            if 'Heading 2' not in doc.styles:
                style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = '黑体'
                style.font.size = Pt(14)
                style.font.bold = True
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                style.paragraph_format.space_after = Pt(10)
        except:
            pass

        try:
            # 标题3
            if 'Heading 3' not in doc.styles:
                style = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = '黑体'
                style.font.size = Pt(12)
                style.font.bold = True
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                style.paragraph_format.space_after = Pt(8)
        except:
            pass

    def parse_markdown_line(self, line: str) -> Dict:
        """解析markdown行"""
        line = line.rstrip()

        # 代码块检测
        if line.strip().startswith('```'):
            if self.in_code_block:
                self.in_code_block = False
                return {'type': 'code_end'}
            else:
                self.in_code_block = True
                return {'type': 'code_start', 'content': line.strip()[3:]}
            return {'type': 'code_block'}

        # 在代码块中，直接返回原始内容
        if self.in_code_block:
            return {'type': 'code_content', 'content': line}

        # 表格检测
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not self.in_table:
                self.in_table = True
                self.table_data = []
                self.table_headers = []

            # 解析表格行
            cells = [cell.strip() for cell in line.strip()[1:-1].split('|')]

            # 检查是否是分隔线
            if all('---' in cell or cell.strip() == '' for cell in cells):
                return {'type': 'table_separator'}
            else:
                # 如果是第一行，认为是表头
                if not self.table_headers and not self.table_data:
                    self.table_headers = cells
                    return {'type': 'table_header', 'cells': cells}
                else:
                    self.table_data.append(cells)
                    return {'type': 'table_row', 'cells': cells}
        else:
            if self.in_table:
                self.in_table = False
                table_info = {
                    'type': 'table_complete',
                    'headers': self.table_headers,
                    'data': self.table_data
                }
                self.table_headers = []
                self.table_data = []
                return table_info

        # 标题检测
        if line.startswith('# '):
            return {'type': 'heading', 'level': 1, 'content': line[2:].strip()}
        elif line.startswith('## '):
            return {'type': 'heading', 'level': 2, 'content': line[3:].strip()}
        elif line.startswith('### '):
            return {'type': 'heading', 'level': 3, 'content': line[4:].strip()}
        elif line.startswith('#### '):
            return {'type': 'heading', 'level': 4, 'content': line[5:].strip()}

        # 列表检测
        if line.startswith('- ') or line.startswith('* '):
            return {'type': 'list_item', 'content': line[2:].strip()}
        elif re.match(r'^\d+\.\s+', line):
            match = re.match(r'^(\d+)\.\s+(.*)', line)
            return {'type': 'list_item', 'content': match.group(2), 'number': match.group(1)}

        # 引用检测
        if line.startswith('> '):
            return {'type': 'quote', 'content': line[2:].strip()}

        # 空行
        if not line.strip():
            return {'type': 'empty'}

        # 普通段落
        return {'type': 'paragraph', 'content': line}

    def add_heading(self, doc: Document, content: str, level: int):
        """添加标题"""
        if level == 1:
            style = 'Heading 1'
        elif level == 2:
            style = 'Heading 2'
        elif level == 3:
            style = 'Heading 3'
        else:
            style = 'Normal'

        paragraph = doc.add_paragraph(content, style=style)
        return paragraph

    def add_paragraph(self, doc: Document, content: str):
        """添加段落"""
        if content.strip():
            paragraph = doc.add_paragraph(content, style='Normal')
            return paragraph

    def add_list_item(self, doc: Document, content: str, level: int = 0):
        """添加列表项"""
        paragraph = doc.add_paragraph(style='Normal')
        run = paragraph.add_run()

        # 添加缩进
        if level > 0:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)

        # 添加项目符号
        run.add_text('• ' if level == 0 else '◦ ')
        run.add_text(content)

    def add_code_block(self, doc: Document, code_lines: List[str], language: str = ''):
        """添加代码块"""
        if not code_lines:
            return

        # 添加代码标题
        if language:
            paragraph = doc.add_paragraph(f"代码块 ({language})", style='Normal')
            paragraph.paragraph_format.space_after = Pt(2)

        # 添加代码内容
        for line in code_lines:
            paragraph = doc.add_paragraph(style='Normal')
            run = paragraph.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(2)

    def add_table(self, doc: Document, headers: List[str], data: List[List[str]]):
        """添加表格"""
        if not headers and not data:
            return

        # 创建表格
        rows = len(data) + 1 if data else 1
        cols = len(headers) if headers else (len(data[0]) if data else 1)

        if cols == 0:
            return

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # 设置表头
        if headers:
            for i, header in enumerate(headers):
                if i < cols:
                    cell = table.cell(0, i)
                    cell.text = header
                    # 设置表头样式
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        # 设置数据行
        for row_idx, row_data in enumerate(data):
            if row_idx + 1 >= rows:
                break
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < cols:
                    cell = table.cell(row_idx + 1, col_idx)
                    cell.text = cell_data

    def add_quote(self, doc: Document, content: str):
        """添加引用"""
        paragraph = doc.add_paragraph(content, style='Normal')
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        # 设置引用样式（灰色）
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)

    def convert_file(self, input_file: str, output_file: str, title: str = None):
        """转换单个文件"""
        if not os.path.exists(input_file):
            print(f"错误: 输入文件不存在: {input_file}")
            return False

        print(f"转换文件: {input_file}")

        # 读取markdown内容
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # 创建Word文档
        self.doc = self.create_word_document()

        # 添加标题
        if title:
            self.add_heading(self.doc, title, 1)
            self.doc.add_paragraph()  # 空行

        # 解析并转换内容
        code_block_lines = []
        current_language = ''

        for line in lines:
            parsed = self.parse_markdown_line(line)

            if parsed['type'] == 'heading':
                self.add_heading(self.doc, parsed['content'], parsed['level'])
                self.doc.add_paragraph()  # 空行

            elif parsed['type'] == 'paragraph':
                self.add_paragraph(self.doc, parsed['content'])

            elif parsed['type'] == 'list_item':
                level = 0  # 可以根据缩进调整
                self.add_list_item(self.doc, parsed['content'], level)

            elif parsed['type'] == 'code_start':
                current_language = parsed['content']
                code_block_lines = []

            elif parsed['type'] == 'code_content':
                code_block_lines.append(parsed['content'])

            elif parsed['type'] == 'code_end':
                self.add_code_block(self.doc, code_block_lines, current_language)
                code_block_lines = []
                current_language = ''
                self.doc.add_paragraph()  # 空行

            elif parsed['type'] == 'table_complete':
                self.add_table(self.doc, parsed['headers'], parsed['data'])
                self.doc.add_paragraph()  # 空行

            elif parsed['type'] == 'quote':
                self.add_quote(self.doc, parsed['content'])

            elif parsed['type'] == 'empty':
                # 空行，添加小间距
                if self.doc.paragraphs and self.doc.paragraphs[-1].text.strip():
                    self.doc.add_paragraph()

        # 保存文档
        self.doc.save(output_file)
        print(f"✅ 已保存: {output_file}")
        return True

    def convert_directory(self, input_dir: str, output_dir: str, file_pattern: str = "*.md"):
        """转换目录中的所有markdown文件"""
        input_path = Path(input_dir)
        output_path = Path(output_dir)

        if not input_path.exists():
            print(f"错误: 输入目录不存在: {input_dir}")
            return

        output_path.mkdir(exist_ok=True)

        # 查找所有markdown文件
        md_files = list(input_path.glob(file_pattern))

        if not md_files:
            print(f"在 {input_dir} 中未找到匹配的markdown文件")
            return

        print(f"找到 {len(md_files)} 个markdown文件，开始转换...")

        success_count = 0
        for md_file in md_files:
            # 生成输出文件名
            output_file = output_path / f"{md_file.stem}.docx"

            # 转换文件
            if self.convert_file(str(md_file), str(output_file), md_file.stem):
                success_count += 1

        print(f"\n转换完成: {success_count}/{len(md_files)} 个文件成功")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="Markdown转Word转换工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 转换单个文件
  python markdown_to_word.py -i input.md -o output.docx

  # 转换目录中的所有文件
  python markdown_to_word.py -d references/ -o output/

  # 转换生成的文档
  python markdown_to_word.py -d generated_docs/ -o word_output/

  # 转换并指定标题
  python markdown_to_word.py -i input.md -o output.docx -t "项目文档"
        """
    )

    parser.add_argument('-i', '--input', help='输入markdown文件路径')
    parser.add_argument('-o', '--output', help='输出Word文件路径或目录')
    parser.add_argument('-d', '--directory', help='输入目录路径（批量转换）')
    parser.add_argument('-t', '--title', help='文档标题')
    parser.add_argument('-p', '--pattern', default='*.md', help='文件匹配模式（批量转换时使用）')

    args = parser.parse_args()

    converter = MarkdownToWordConverter()

    if args.input and args.output:
        # 单个文件转换
        title = args.title or Path(args.input).stem
        converter.convert_file(args.input, args.output, title)

    elif args.directory and args.output:
        # 批量转换
        converter.convert_directory(args.directory, args.output, args.pattern)

    else:
        print("错误: 请指定输入和输出")
        print("使用 -h 查看帮助")
        exit(1)


if __name__ == '__main__':
    main()